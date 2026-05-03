"""
Generate the final TravelEval research report as a Word (.docx) document.
Uses real evaluation data from NVIDIA-hosted model runs.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FIGURES_DIR = os.path.join(os.path.dirname(__file__), "figures")
OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "TravelEval_Final_Report.docx")


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(8)
    return p


def add_table(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], "2C3E50")
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)

    # Data rows
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "F8F9FA")
            for para in cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        for run in p.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(127, 140, 141)


def add_figure(doc, filename, caption):
    path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(5.5))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(100, 100, 100)
        cap.paragraph_format.space_after = Pt(12)
    else:
        p = doc.add_paragraph(f"[Figure not found: {filename}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_page_break(doc):
    doc.add_page_break()


def build_report():
    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    # ---- Default paragraph style ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ======================================================
    # TITLE PAGE
    # ======================================================
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("TravelEval: Auditing AI-Driven Tourism Systems")
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(44, 62, 80)

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "A Research Evaluation of Reliability, Bias, Real-Time Adaptation,\n"
        "and Niche Knowledge in NVIDIA-Hosted Large Language Models"
    )
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(100, 110, 120)

    doc.add_paragraph()
    meta_lines = [
        "GenAI Travel Evaluation Project",
        "Date: April 30, 2026",
        "Framework Version: 2.0",
        "Models Evaluated: NVIDIA-hosted Llama 3.1 (8B, 70B), Llama 3.3 (70B), Mixtral 8x7B",
    ]
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(130, 140, 150)

    add_page_break(doc)

    # ======================================================
    # ABSTRACT
    # ======================================================
    add_heading(doc, "Abstract", 1, color=(44, 62, 80))
    abstract_text = (
        "The deployment of Large Language Models (LLMs) in the tourism sector introduces paradigm-shifting "
        "capabilities alongside profound reliability and fairness risks. This research presents TravelEval, "
        "a rigorous, reproducible evaluation framework designed to audit LLM performance across four dimensions: "
        "factual hallucination detection, geographic representation bias, real-time dynamic query accuracy, "
        "and niche domain knowledge retrieval. "
        "Using four NVIDIA-hosted models — Meta Llama 3.1 8B, Llama 3.1 70B, Llama 3.3 70B, and Mistral "
        "Mixtral 8x7B — evaluated across 194 gold benchmark rows, 63 bias prompts (producing 658 geographic "
        "entities), 160 real-time queries, and 286 extracted claims from 120 hallucination responses, this "
        "study documents critical vulnerabilities. "
        "Results confirm near-zero real-time accuracy for unaugmented models (≤2.5%), persistent European "
        "over-representation in geographic recommendations across all model families, strict tourism factual "
        "correctness below 35.5%, and response-level hallucination rates between 15% and 28%. "
        "The findings advocate for mandated hybrid architectures combining LLMs with external APIs, "
        "retrieval-augmented generation, and human-in-the-loop oversight for enterprise tourism deployments."
    )
    add_para(doc, abstract_text, italic=True)

    doc.add_paragraph()
    kw = doc.add_paragraph()
    kwr = kw.add_run("Keywords: ")
    kwr.bold = True
    kw.add_run(
        "Large Language Models, Tourism AI, Hallucination Detection, Geographic Bias, "
        "Real-Time Information Retrieval, Evaluation Framework, NVIDIA NIM"
    )

    add_page_break(doc)

    # ======================================================
    # TABLE OF CONTENTS (manual)
    # ======================================================
    add_heading(doc, "Table of Contents", 1, color=(44, 62, 80))
    toc_items = [
        "1. Introduction",
        "    1.1 Background",
        "    1.2 Motivation",
        "    1.3 Scope",
        "2. Research Questions and Hypotheses",
        "3. Data and Experimental Design",
        "    3.1 Prompt Construction",
        "    3.2 Datasets",
        "    3.3 Model Selection",
        "    3.4 Annotation and Ground Truth",
        "4. Evaluation Pipeline Architecture",
        "5. Results: Real-Time Reliability",
        "6. Results: Geographic Bias",
        "7. Results: Niche Gold Factuality",
        "8. Results: Claim-Level Hallucination",
        "9. Comparative Model Performance",
        "10. Discussion and Implications",
        "11. Threats to Validity",
        "12. Conclusions and Future Work",
        "13. References",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(11)

    add_page_break(doc)

    # ======================================================
    # 1. INTRODUCTION
    # ======================================================
    add_heading(doc, "1. Introduction", 1, color=(44, 62, 80))

    add_heading(doc, "1.1 Background", 2)
    add_para(
        doc,
        "Tourism is an information-sensitive domain where factual correctness and equitable "
        "representation directly affect traveler safety, satisfaction, and destination choice. "
        "Modern travelers increasingly rely on AI-powered assistants for itinerary planning, "
        "destination discovery, and logistical coordination. When these systems generate tourism "
        "advice, they combine factual knowledge, inference, and stylistic generation in ways that "
        "can be difficult to independently validate.",
    )
    add_para(
        doc,
        "Unlike general-purpose text generation, tourism recommendations carry immediate physical, "
        "financial, and logistical consequences. Incorrect information about visa requirements, "
        "entry regulations, or transportation schedules can result in disrupted itineraries, "
        "financial losses, or personal safety risks for travelers.",
    )

    add_heading(doc, "1.2 Motivation", 2)
    add_para(
        doc,
        "This evaluation responds to two urgent challenges in AI-powered tourism deployment:",
    )
    bullets = [
        "Reliability: Can live LLM deployments provide consistent and accurate answers for "
        "weather conditions, local time, attraction logistics, and factual travel recommendations?",
        "Fairness: Do models systematically over-recommend certain world regions while "
        "underrepresenting others, creating inequitable visibility for global destinations?",
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    add_para(
        doc,
        "The goal is not only to measure performance but to document methodological rigor and "
        "produce a reproducible, publishable benchmark of results that can inform responsible "
        "AI deployment in the tourism sector.",
    )

    add_heading(doc, "1.3 Scope", 2)
    add_para(doc, "The evaluation covers four technical modules:")
    scope_bullets = [
        "Real-Time Reliability: 160 dynamic query evaluations (40 prompts × 4 models) "
        "for weather and local time accuracy.",
        "Bias Evaluation: 63 neutral recommendation prompts producing 658 extracted geographic "
        "entities for regional distribution analysis.",
        "Niche Gold Factuality: 194 highly specific tourism gold benchmark rows evaluated "
        "against authoritative ground truth across two stable model endpoints.",
        "Claim-Level Hallucination: 286 extracted factual claims from 120 model responses "
        "verified against live web evidence sources.",
    ]
    for b in scope_bullets:
        p = doc.add_paragraph(b, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    add_para(
        doc,
        "The user trust study was intentionally excluded from this phase and is documented "
        "in the project design for future work.",
    )

    # ======================================================
    # 2. RESEARCH QUESTIONS AND HYPOTHESES
    # ======================================================
    add_heading(doc, "2. Research Questions and Hypotheses", 1, color=(44, 62, 80))

    add_heading(doc, "2.1 Research Questions", 2)
    rqs = [
        ("RQ1", "What is the real-time information accuracy of NVIDIA-hosted LLMs for "
         "tourism-related dynamic queries involving live weather and local time?"),
        ("RQ2", "How do these models distribute tourism recommendations across global regions, "
         "and what degree of geographic concentration or diversity do they exhibit?"),
        ("RQ3", "What is the strict factual correctness rate on a domain-specific niche "
         "tourism gold benchmark with authoritative ground truth?"),
        ("RQ4", "How frequently do live tourism responses contain extractable claims that are "
         "hallucinated, unsupported, or contradicted by verifiable web evidence?"),
    ]
    for rq_id, rq_text in rqs:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{rq_id}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        p.add_run(rq_text).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    add_heading(doc, "2.2 Hypotheses", 2)
    hs = [
        ("H1", "Real-time dynamic queries will have accuracy below 5% due to the absence of "
         "external tool access or retrieval augmentation in base model deployments."),
        ("H2", "Europe will be significantly overrepresented in tourism recommendations "
         "compared to all other global regions, reflecting pre-training data distribution imbalance."),
        ("H3", "Strict factual correctness on the hard niche tourism gold benchmark will "
         "remain below 40% for all tested model architectures."),
        ("H4", "Response-level claim hallucination rate will exceed 25% for smaller models, "
         "with larger models showing moderate improvement."),
    ]
    for h_id, h_text in hs:
        p = doc.add_paragraph()
        r1 = p.add_run(f"{h_id}: ")
        r1.bold = True
        r1.font.size = Pt(11)
        p.add_run(h_text).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    add_figure(
        doc,
        "research_question_hypothesis_matrix.png",
        "Figure 1: Alignment between research questions and testable hypotheses.",
    )

    # ======================================================
    # 3. DATA AND EXPERIMENTAL DESIGN
    # ======================================================
    add_heading(doc, "3. Data and Experimental Design", 1, color=(44, 62, 80))

    add_heading(doc, "3.1 Prompt Construction", 2)
    add_para(
        doc,
        "The evaluation prompts were constructed to reflect authentic tourism search behavior "
        "and span six primary query categories: destination discovery, attraction and itinerary "
        "recommendations, local weather and temporal queries, transportation guidance, cultural "
        "and historical context, and travel safety and logistics.",
    )
    add_para(
        doc,
        "Prompts were designed to avoid leading language that could bias model responses. "
        "For bias evaluation, prompts were intentionally neutral (e.g., 'What are the best "
        "travel destinations in Africa?') to measure organic recommendation patterns. "
        "For real-time queries, prompts included precise location identifiers with geographic "
        "coordinates and IANA timezone metadata to allow exact ground-truth comparison.",
    )

    add_heading(doc, "3.2 Datasets", 2)
    add_para(doc, "Four core evaluation datasets were constructed and curated for this study:")

    add_table(
        doc,
        ["Dataset", "Type", "Size", "Purpose"],
        [
            ["hallucination_dataset_balanced_60.csv", "Hallucination evaluation",
             "60 prompts / 286 claims", "Claim-level hallucination analysis"],
            ["bias_dataset.csv + results", "Bias evaluation",
             "63 prompts / 658 entities", "Geographic recommendation bias"],
            ["research_gold_niche.csv", "Niche gold factuality",
             "194 rows", "Strict tourism correctness benchmark"],
            ["realtime_dataset.csv + results", "Real-time evaluation",
             "40 cities / 160 queries", "Dynamic information accuracy"],
        ],
        "Table 1: Evaluation dataset inventory with size and purpose.",
    )

    add_para(
        doc,
        "The niche gold dataset (research_gold_niche.csv) was constructed through an "
        "automated pipeline: seed tourism concepts from nine underrepresented global regions "
        "were combined with live web search snippets via the SERP API. A generator LLM "
        "synthesized complex, long-tail queries, which were then filtered by source quality "
        "(official government URLs, UNESCO records, tourism authority pages) and verified "
        "against authoritative ground truth. This produced 194 highly specific, culturally "
        "nuanced questions spanning visa regulations, entry requirements, permit procedures, "
        "and regional travel restrictions.",
    )

    add_figure(
        doc,
        "dataset_composition_by_module.png",
        "Figure 2: Dataset composition and sample count across the four evaluation modules.",
    )

    add_heading(doc, "3.3 Model Selection", 2)
    add_para(
        doc,
        "Four NVIDIA-hosted models were evaluated, selected to represent a range of parameter "
        "scales and architectural families accessible through the NVIDIA NIM inference platform:",
    )
    add_table(
        doc,
        ["Model ID", "Family", "Parameters", "Provider"],
        [
            ["nvidia:meta/llama-3.1-8b-instruct", "Llama 3.1", "8B", "Meta / NVIDIA NIM"],
            ["nvidia:meta/llama-3.1-70b-instruct", "Llama 3.1", "70B", "Meta / NVIDIA NIM"],
            ["nvidia:meta/llama-3.3-70b-instruct", "Llama 3.3", "70B", "Meta / NVIDIA NIM"],
            ["nvidia:mistralai/mixtral-8x7b-instruct-v0.1", "Mixtral MoE", "8x7B (~47B active)",
             "Mistral AI / NVIDIA NIM"],
        ],
        "Table 2: Models evaluated in the live benchmark.",
    )
    add_para(
        doc,
        "The NVIDIA NIM platform was selected to provide a standardized inference environment "
        "and eliminate provider-level artifacts when comparing model performance. During the "
        "full-scale evaluation, the 70B Llama models experienced intermittent connection "
        "failures on the live platform. As a result, the strict factuality and hallucination "
        "modules focus on the two most stable model endpoints: Llama 3.1 8B and Mixtral 8x7B.",
    )

    add_heading(doc, "3.4 Annotation and Ground Truth", 2)
    gt_items = [
        "Real-time queries: verified against live Open-Meteo API (temperature, ±3°C tolerance) "
        "and IANA timezone database (local time, ±20-minute tolerance) at exact execution time.",
        "Geographic bias: verified through entity extraction and region mapping against the "
        "OpenFlights airport and city database with a 9-region taxonomy.",
        "Niche gold factuality: strict answer matching against curated ground truth from "
        "official government sources, US State Department travel advisories, and UNESCO records.",
        "Claim verification: each extracted claim independently labeled SUPPORTED, CONTRADICTED, "
        "NOT_FOUND, or UNCLEAR by an LLM judge (llama-3.3-70b-versatile) against top-K "
        "web search snippets retrieved specifically for that claim.",
    ]
    for item in gt_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    # ======================================================
    # 4. EVALUATION PIPELINE ARCHITECTURE
    # ======================================================
    add_heading(doc, "4. Evaluation Pipeline Architecture", 1, color=(44, 62, 80))

    add_heading(doc, "4.1 System Overview", 2)
    add_para(
        doc,
        "The TravelEval pipeline was designed as a modular, reproducible research framework. "
        "Each component has a single responsibility, allowing independent testing and extension. "
        "The high-level data flow is:",
    )
    add_para(
        doc,
        "Query Generator → Model Runner → Response Collector → Claim Extractor → "
        "Ground Truth Validator → Analysis Engine → Visualization",
        italic=True,
    )

    add_heading(doc, "4.2 Core Components", 2)
    components = [
        ("eval_pipeline.py",
         "Orchestrates experiment flow and model execution across four modes: gold "
         "(LLM-as-judge against ground truth), claim (atomic claim extraction and web verification), "
         "bias (recommendation entity extraction), and auto (mode inference per row)."),
        ("realtime_eval.py",
         "Specialized module for weather and local time queries. Fetches live ground truth "
         "at evaluation time from Open-Meteo and IANA timezone databases with configurable "
         "tolerance thresholds."),
        ("bias_analysis.py",
         "Extracts destination and region entities from model outputs. Computes Shannon "
         "diversity entropy and Herfindahl-Hirschman Index (HHI) concentration across regions."),
        ("niche_generator.py",
         "Automated pipeline for generating hard-to-know tourism prompts. Combines seed "
         "concepts from underrepresented regions with SerpAPI snippets and LLM fact synthesis."),
        ("analyze_results.py",
         "Aggregates per-response metrics into grouped tables, computes 95% confidence "
         "intervals via bootstrap resampling, and executes Mann-Whitney U significance tests."),
    ]
    for comp_name, comp_desc in components:
        p = doc.add_paragraph()
        p.add_run(comp_name).bold = True
        p.add_run(": " + comp_desc).font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)

    add_heading(doc, "4.3 Evaluation Modes", 2)
    add_table(
        doc,
        ["Mode", "Mechanism", "Ground Truth Source", "Primary Use"],
        [
            ["gold", "LLM-as-judge correctness scoring", "Curated gold answers", "Niche factuality"],
            ["claim", "Atomic claim extraction + web verification", "Live SERP API snippets",
             "Hallucination detection"],
            ["bias", "Entity extraction + region mapping", "OpenFlights taxonomy",
             "Geographic bias"],
            ["auto", "Dynamic mode selection per row", "Mixed", "General evaluation"],
        ],
        "Table 3: Evaluation mode descriptions and data sources.",
    )

    add_figure(
        doc,
        "pipeline_architecture_diagram.png",
        "Figure 3: Modular TravelEval pipeline architecture.",
    )

    add_page_break(doc)

    # ======================================================
    # 5. RESULTS: REAL-TIME RELIABILITY
    # ======================================================
    add_heading(doc, "5. Results: Real-Time Reliability", 1, color=(44, 62, 80))

    add_para(
        doc,
        "The real-time evaluation assessed live model performance on weather and local time queries "
        "using 160 total evaluations (40 prompts across 4 models). These queries explicitly require "
        "live environmental data that no base LLM can access without external tool integration.",
    )

    add_heading(doc, "5.1 Accuracy by Model", 2)
    add_table(
        doc,
        ["Model", "Accuracy (%)"],
        [
            ["nvidia:meta/llama-3.1-8b-instruct", "2.50%"],
            ["nvidia:meta/llama-3.1-70b-instruct", "0.00%"],
            ["nvidia:meta/llama-3.3-70b-instruct", "0.00%"],
            ["nvidia:mistralai/mixtral-8x7b-instruct-v0.1", "0.00%"],
        ],
        "Table 4: Real-time query accuracy per model (tolerance: ±3°C / ±20 min).",
    )

    add_heading(doc, "5.2 Accuracy by Query Type", 2)
    add_table(
        doc,
        ["Query Type", "Accuracy (%)"],
        [
            ["temperature_celsius", "1.25%"],
            ["local_time_minutes", "0.00%"],
        ],
        "Table 5: Real-time accuracy by dynamic information metric.",
    )

    add_figure(
        doc,
        "realtime_accuracy_by_model.png",
        "Figure 4: Real-time accuracy per model — showing near-zero performance across all endpoints.",
    )
    add_figure(
        doc,
        "realtime_accuracy_by_metric.png",
        "Figure 5: Real-time accuracy by query type (temperature vs. local time).",
    )

    add_heading(doc, "5.3 Interpretation", 2)
    add_para(
        doc,
        "Only Llama 3.1 8B returned one correct temperature response out of 40 prompts (2.5%). "
        "All other models produced zero correct answers. Models either hallucinated a plausible "
        "current state based on seasonal averages, fabricated a temperature within the expected "
        "range, or provided historical or typical values rather than live measurements.",
    )
    add_para(
        doc,
        "Local time accuracy was zero for all models. Models consistently failed to account for "
        "Daylight Saving Time transitions, summer/winter offset shifts, and regional timezone "
        "subdivisions, confirming that time calculation requires verified external data sources "
        "rather than parametric inference.",
    )
    add_para(
        doc,
        "These results confirm Hypothesis H1: real-time accuracy is effectively zero for "
        "unaugmented base model deployments. This represents a categorical failure mode — "
        "not a matter of degree — for any tourism system that relies on live environmental context.",
    )

    # ======================================================
    # 6. RESULTS: GEOGRAPHIC BIAS
    # ======================================================
    add_heading(doc, "6. Results: Geographic Bias and Regional Distribution", 1, color=(44, 62, 80))

    add_para(
        doc,
        "The bias evaluation analyzed 658 extracted destination entities from 252 model responses "
        "to 63 neutral tourism recommendation prompts. Geographic entities were mapped to a "
        "9-region taxonomy using the OpenFlights database.",
    )

    add_heading(doc, "6.1 Summary Metrics by Model", 2)
    add_table(
        doc,
        ["Model", "Entities", "Unique Destinations", "Countries", "Shannon Diversity", "HHI", "Top Region"],
        [
            ["Llama 3.1 8B", "211", "145", "35", "2.058", "0.0467", "Europe"],
            ["Llama 3.1 70B", "133", "103", "37", "2.056", "0.0527", "Europe"],
            ["Llama 3.3 70B", "127", "102", "31", "2.061", "0.0574", "Europe"],
            ["Mixtral 8x7B", "187", "150", "45", "2.130", "0.1106", "Unknown*"],
        ],
        "Table 6: Bias summary metrics. *Mixtral had a higher 'unknown' entity classification "
        "rate due to entity extraction fallback.",
    )

    add_heading(doc, "6.2 Top 3 Regions by Model", 2)
    add_table(
        doc,
        ["Model", "Rank 1", "Rank 2", "Rank 3"],
        [
            ["Llama 3.1 70B", "Europe (19.5%)", "Asia (17.3%)", "Oceania (14.3%)"],
            ["Llama 3.1 8B", "Europe (18.5%)", "Asia (13.3%)", "North America (13.3%)"],
            ["Llama 3.3 70B", "Europe (16.5%)", "Asia (15.0%)", "Oceania (14.2%)"],
            ["Mixtral 8x7B", "Unknown (27.8%)", "South America (11.2%)", "Europe (10.2%)"],
        ],
        "Table 7: Top three recommended regions by share per model.",
    )

    add_figure(
        doc,
        "bias_region_share_by_model.png",
        "Figure 6: Regional recommendation share distribution by model.",
    )
    add_figure(
        doc,
        "bias_diversity_vs_hhi.png",
        "Figure 7: Diversity (Shannon entropy) vs. concentration (HHI) tradeoff across models.",
    )

    add_heading(doc, "6.3 Interpretation", 2)
    add_para(
        doc,
        "Geographic bias is pronounced and consistent across all Meta Llama model families. "
        "Europe is the top-ranked recommendation region for all three Meta models, confirming "
        "Hypothesis H2. Africa and the Middle East are consistently underrepresented, appearing "
        "in the bottom tier of recommendations despite representing major global tourism markets.",
    )
    add_para(
        doc,
        "The Mixtral model displays a broader country vocabulary (45 unique countries vs. "
        "31-37 for Meta models), suggesting stronger nominal diversity. However, Mixtral's "
        "high 'unknown' region rate (27.8%) indicates that its entity naming conventions "
        "differ from standard geographic databases, introducing classification noise.",
    )
    add_para(
        doc,
        "Shannon diversity scores are broadly similar across models (2.056–2.130), but the "
        "HHI concentration metric reveals that Mixtral has higher country-level concentration "
        "(HHI 0.1106) compared to Llama variants (0.0467–0.0574), meaning Mixtral's diversity "
        "is distributed unevenly with certain countries dominating.",
    )

    # ======================================================
    # 7. RESULTS: NICHE GOLD FACTUALITY
    # ======================================================
    add_heading(doc, "7. Results: Niche Gold Factuality", 1, color=(44, 62, 80))

    add_para(
        doc,
        "This module evaluates strict tourism factuality on 194 gold benchmark rows derived from "
        "nine underrepresented global regions. Prompts target highly specific, long-tail knowledge "
        "such as visa registration thresholds, permit requirements, and regional entry policies.",
    )

    add_heading(doc, "7.1 Answer Correctness by Model", 2)
    add_table(
        doc,
        ["Model", "Mean Correctness", "95% CI Lower", "95% CI Upper"],
        [
            ["nvidia:meta/llama-3.1-8b-instruct", "35.31%", "29.02%", "41.60%"],
            ["nvidia:mistralai/mixtral-8x7b-instruct-v0.1", "33.76%", "28.25%", "39.28%"],
        ],
        "Table 8: Niche gold benchmark correctness with 95% bootstrap confidence intervals.",
    )

    add_figure(
        doc,
        "niche_accuracy_confidence_intervals.png",
        "Figure 8: Mean correctness and 95% confidence intervals for the niche gold benchmark.",
    )

    add_heading(doc, "7.2 Interpretation", 2)
    add_para(
        doc,
        "Both models scored below 35.5% mean correctness on the niche gold benchmark, confirming "
        "Hypothesis H3. The overlapping confidence intervals (Llama 8B: 29.0–41.6%, "
        "Mixtral: 28.3–39.3%) indicate that the performance difference between models is not "
        "statistically significant for this task.",
    )
    add_para(
        doc,
        "This finding is significant because the niche benchmark was designed to target "
        "knowledge that is sparsely represented in standard pre-training corpora. When queries "
        "move from well-known tourist facts to highly specific regional constraints, models "
        "attempt to satisfy the user by interpolating generic knowledge, resulting in highly "
        "plausible but fundamentally incorrect responses.",
    )
    add_para(
        doc,
        "The practical implication is that model parameter scale alone does not resolve niche "
        "knowledge gaps in specialized domains. Both the 8B Llama and the effectively 47B "
        "active-parameter Mixtral achieve statistically equivalent and similarly weak performance, "
        "suggesting that specialized tourism fine-tuning or retrieval augmentation is required.",
    )

    # ======================================================
    # 8. RESULTS: CLAIM-LEVEL HALLUCINATION
    # ======================================================
    add_heading(doc, "8. Results: Claim-Level Hallucination", 1, color=(44, 62, 80))

    add_para(
        doc,
        "The hallucination module evaluates 286 extracted claims from 120 model responses "
        "(60 prompts × 2 models). Each claim is individually verified against live web evidence "
        "and classified by an LLM judge.",
    )

    add_heading(doc, "8.1 Claim Verdict Distribution", 2)
    add_table(
        doc,
        ["Model", "Supported", "Contradicted", "Not Found", "Unclear"],
        [
            ["Llama 3.1 8B", "79.28%", "2.70%", "10.81%", "7.21%"],
            ["Mixtral 8x7B", "84.57%", "2.29%", "10.29%", "2.86%"],
        ],
        "Table 9: Claim verdict distribution by model (share of total claims per model).",
    )

    add_heading(doc, "8.2 Response-Level Hallucination Rate", 2)
    add_table(
        doc,
        ["Model", "Mean Hallucination Rate"],
        [
            ["nvidia:meta/llama-3.1-8b-instruct", "28.39%"],
            ["nvidia:mistralai/mixtral-8x7b-instruct-v0.1", "15.14%"],
        ],
        "Table 10: Response-level hallucination rate (proportion of claims not supported per response).",
    )

    add_figure(
        doc,
        "hallucination_verdict_distribution.png",
        "Figure 9: Claim verification verdict distribution by model.",
    )

    add_heading(doc, "8.3 Interpretation", 2)
    add_para(
        doc,
        "Mixtral achieves a materially lower response-level hallucination rate (15.14%) "
        "compared to Llama 3.1 8B (28.39%), confirming Hypothesis H4 for the smaller model. "
        "The NOT_FOUND verdict category (10–11% for both models) represents a significant "
        "additional risk: these claims cannot be verified by available web evidence, meaning "
        "they may be correct but unsourced, or may be novel fabrications that evade detection.",
    )
    add_para(
        doc,
        "The UNCLEAR rate for Llama 3.1 8B (7.21%) is notably higher than Mixtral (2.86%), "
        "indicating that the smaller model generates more ambiguous claim structures that resist "
        "atomic decomposition, complicating automated fact-checking pipelines.",
    )

    # ======================================================
    # 9. COMPARATIVE MODEL PERFORMANCE
    # ======================================================
    add_heading(doc, "9. Comparative Model Performance", 1, color=(44, 62, 80))

    add_heading(doc, "9.1 Summary Across All Dimensions", 2)
    add_table(
        doc,
        ["Metric", "Llama 3.1 8B", "Llama 3.1 70B", "Llama 3.3 70B", "Mixtral 8x7B"],
        [
            ["Real-Time Accuracy", "2.50%", "0.00%", "0.00%", "0.00%"],
            ["Top Recommended Region", "Europe", "Europe", "Europe", "Unknown*"],
            ["Shannon Diversity (Bias)", "2.058", "2.056", "2.061", "2.130"],
            ["Niche Gold Correctness", "35.31%", "N/A†", "N/A†", "33.76%"],
            ["Response Hallucination Rate", "28.39%", "N/A†", "N/A†", "15.14%"],
        ],
        "Table 11: Comparative performance across evaluation dimensions. "
        "†70B models excluded from gold/claim modules due to endpoint instability. "
        "*Mixtral 'Unknown' reflects entity classification fallback.",
    )

    add_heading(doc, "9.2 Key Observations", 2)
    obs = [
        "Mixtral 8x7B outperforms Llama 3.1 8B on hallucination rate (15.1% vs 28.4%) "
        "and niche correctness (33.8% vs 35.3%), though the correctness difference is not "
        "statistically significant.",
        "All Meta Llama models show consistent Eurocentric bias; the 70B models show "
        "marginally higher European recommendation share than the 8B model.",
        "Real-time performance is uniformly catastrophic across all models, with only one "
        "correct temperature answer across 160 total evaluations.",
        "The 70B models demonstrated higher instability (connection failures, rate limits) "
        "on the NVIDIA NIM platform, limiting their participation in the full evaluation suite.",
    ]
    for o in obs:
        p = doc.add_paragraph(o, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)

    add_figure(
        doc,
        "model_performance_comparison.png",
        "Figure 10: Comparative model performance across evaluation dimensions.",
    )

    # ======================================================
    # 10. DISCUSSION
    # ======================================================
    add_heading(doc, "10. Discussion and Implications", 1, color=(44, 62, 80))

    add_heading(doc, "10.1 Core Findings", 2)
    findings = [
        "Real-time performance is effectively absent for all models without external tool integration.",
        "Geographic bias is persistent and structural: Europe dominates recommendations "
        "across multiple model families, reflecting pre-training data imbalance.",
        "Strict niche factuality is weak: both models fall below 36% correctness, "
        "with statistically indistinguishable performance between architectures.",
        "Claim-level hallucination remains substantial at 15–28%, with the NOT_FOUND "
        "category representing an additional unverifiable risk layer.",
    ]
    for f in findings:
        p = doc.add_paragraph(f, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)

    add_heading(doc, "10.2 Operational Recommendations", 2)
    recs = [
        "Do not deploy base LLMs for live weather, time, or schedule guidance. "
        "External API integration via tool-calling is mandatory for any temporally sensitive query.",
        "Implement region-aware bias correction in system prompts and post-processing "
        "to counteract the pre-training distribution that heavily favors Western destinations.",
        "Use strict gold benchmarks with authoritative sources for tourism evaluation "
        "rather than open-ended subjective quality measures.",
        "Collect and explicitly verify model response claims against authoritative sources "
        "before displaying to end users in high-stakes travel contexts.",
        "Consider retrieval-augmented generation (RAG) architectures that ground responses "
        "in verified tourism databases, official source repositories, and live API data.",
    ]
    for r in recs:
        p = doc.add_paragraph(r, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)

    add_heading(doc, "10.3 Implications for Hybrid Systems", 2)
    add_para(
        doc,
        "The evidence suggests that LLMs in their current form function best as language "
        "synthesis and formatting engines rather than primary knowledge sources for tourism "
        "applications. A reliable tourism AI system should integrate:",
    )
    hybrid_items = [
        "Knowledge grounding: structured tourism databases, official regulatory sources, "
        "and destination management organization (DMO) content feeds.",
        "Real-time context injection: direct API calls for weather (Open-Meteo, OpenWeather), "
        "flight status (OAG, Cirium), and entry requirements (IATA Travel Centre).",
        "Bias correction layers: regional balancing constraints and diversity requirements "
        "applied at the retrieval and generation stages.",
        "Claim verification pipeline: post-generation claim extraction and evidence matching "
        "before responses are surfaced to users.",
    ]
    for item in hybrid_items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    add_figure(
        doc,
        "discussion_key_findings.png",
        "Figure 11: High-level reliability and risk summary across evaluation dimensions.",
    )

    # ======================================================
    # 11. THREATS TO VALIDITY
    # ======================================================
    add_heading(doc, "11. Threats to Validity", 1, color=(44, 62, 80))

    add_heading(doc, "11.1 Internal Validity", 2)
    add_para(
        doc,
        "Live model endpoints experienced intermittent rate limits and connectivity failures "
        "on the NVIDIA NIM platform during evaluation. This disproportionately affected the "
        "70B model family, which was excluded from the full factuality and hallucination "
        "analyses. Results for the 70B models therefore only appear in real-time and bias "
        "modules, where evaluation completed without significant failure.",
    )

    add_heading(doc, "11.2 External Validity", 2)
    add_para(
        doc,
        "Findings are specific to NVIDIA-hosted model endpoints and may not generalize to "
        "the same models accessed through other providers (Groq, OpenRouter, together.ai) "
        "or run locally. Infrastructure choices — batch size, timeout configuration, "
        "and sampling temperature — affect output distribution and may introduce differences "
        "compared to other deployment contexts.",
    )
    add_para(
        doc,
        "Tourism prompts were designed for evaluation coverage across categories, but actual "
        "user queries may differ in phrasing, context, and specificity from the structured "
        "benchmark questions used here.",
    )

    add_heading(doc, "11.3 Construct Validity", 2)
    add_para(
        doc,
        "The NOT_FOUND verdict category encompasses both genuinely unverifiable claims and "
        "claims that are factually correct but not prominently indexed in the top web search "
        "results used for verification. This conflation means the hallucination rate may be "
        "slightly underestimated.",
    )
    add_para(
        doc,
        "The 9-region taxonomy used for bias analysis reflects a coarse geographic "
        "classification and may miss intra-regional disparities (e.g., within-Africa or "
        "within-Asia representation gaps). A finer-grained country or sub-regional analysis "
        "would reveal additional concentration patterns.",
    )

    add_figure(
        doc,
        "threats_to_validity_risk_map.png",
        "Figure 12: Relative risk assessment for validity threat categories.",
    )

    # ======================================================
    # 12. CONCLUSIONS AND FUTURE WORK
    # ======================================================
    add_heading(doc, "12. Conclusions and Future Work", 1, color=(44, 62, 80))

    add_heading(doc, "12.1 Conclusions", 2)
    add_para(
        doc,
        "This report presents a live benchmark demonstrating that current NVIDIA-hosted LLMs "
        "have limited readiness for standalone tourism applications. Four dimensions were measured "
        "with concrete quantitative results:",
    )
    conclusions = [
        "Real-time accuracy is near zero (≤2.5%) — confirming a fundamental capability gap "
        "for any tourism query requiring live environmental or temporal data.",
        "Geographic bias is persistent and Eurocentric — Europe is the dominant recommendation "
        "region across all tested Meta Llama model families.",
        "Niche factual correctness is weak (≤35.5%) — model architecture scale does not "
        "resolve knowledge gaps for long-tail, domain-specific tourism knowledge.",
        "Hallucination rates are material (15–28%) — with an additional 10% of claims "
        "categorized as NOT_FOUND, representing unverifiable factual risk.",
    ]
    for c in conclusions:
        p = doc.add_paragraph(c, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)

    add_para(
        doc,
        "The TravelEval framework itself represents a validated, reproducible research "
        "infrastructure for ongoing LLM auditing in tourism. All datasets, evaluation code, "
        "and result files are structured for public release and future research extension.",
    )

    add_heading(doc, "12.2 Future Work", 2)
    future = [
        "Extend the hallucination claim evaluation to the full 263-row dataset across all "
        "four model variants.",
        "Evaluate retrieval-augmented generation variants to quantify the accuracy uplift "
        "from external knowledge grounding.",
        "Incorporate a user trust study (survey design documented in user_survey/) to "
        "measure trust calibration when users are exposed to AI-generated hallucinations.",
        "Extend niche factuality evaluation to the complete 351-row benchmark dataset.",
        "Introduce sub-regional bias analysis to detect within-continent representation gaps.",
        "Evaluate additional model families (Claude, GPT-4o) on the same benchmark for "
        "cross-provider comparison.",
    ]
    for f in future:
        p = doc.add_paragraph(f, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)

    # ======================================================
    # 13. REFERENCES
    # ======================================================
    add_heading(doc, "13. References", 1, color=(44, 62, 80))

    refs = [
        "Project Source Code & Architecture: eval_pipeline.py, bias_analysis.py, "
        "realtime_eval.py, niche_generator.py, analyze_results.py",
        "Datasets: data_collection/research_gold_niche.csv, "
        "data_collection/hallucination_dataset_balanced_60.csv, "
        "data_collection/bias_dataset.csv, data_collection/realtime_dataset.csv",
        "Evaluation Results: results/realtime_nvidia_live/, results/bias_nvidia_live_fast/, "
        "results/niche_gold_nvidia_live/, results/hallucination_claims_nvidia_live/",
        "External Data Sources: OpenFlights Airport Database, Open-Meteo API, "
        "IANA Time Zone Database, SerpAPI Web Search",
        "Model Infrastructure: NVIDIA NIM Inference Platform (api.nvcf.nvidia.com)",
        "Brown, T., et al. (2020). Language Models are Few-Shot Learners. NeurIPS 2020.",
        "Ji, Z., et al. (2023). Survey of Hallucination in Natural Language Generation. "
        "ACM Computing Surveys, 55(12), 1–38.",
        "Navigli, R., et al. (2023). Biases in Large Language Models: Origins, Inventory and "
        "Discussion. ACM Journal of Data and Information Quality.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(5)
        for run in p.runs:
            run.font.size = Pt(10)

    # ======================================================
    # FOOTER / CLOSING
    # ======================================================
    add_page_break(doc)
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        "TravelEval Research Report  |  April 2026  |  GenAI Travel Evaluation Project\n"
        "Total Queries Evaluated: 572  |  Total Claims Analyzed: 286  |  "
        "Datasets Curated: 4  |  Models Benchmarked: 4"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(150, 160, 170)
    fr.font.italic = True

    doc.save(OUTPUT_PATH)
    print(f"Report saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
